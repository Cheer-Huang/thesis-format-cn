#!/usr/bin/env python3
"""
清华大学研究生学位论文 DOCX 格式化脚本
按照《清华大学研究生学位论文写作指南》规范修正 Word 文档格式。

用法:
    python format_thu_thesis.py <input.docx> [output.docx]

如不指定 output，则覆盖原文件。

格式规范详见 references/thu_format_spec.md
"""

import sys
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ============================================================
# 工具函数
# ============================================================

def set_run_font(run, cn_font, en_font, size_pt, bold=False, italic=False):
    """设置 run 的中英文字体、字号、粗体、斜体，颜色强制黑色"""
    run.font.size = Pt(size_pt)
    run.font.name = en_font
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), cn_font)
    rFonts.set(qn('w:ascii'), en_font)
    rFonts.set(qn('w:hAnsi'), en_font)


def set_para_spacing_xml(para, before_twips, after_twips, line_twips, line_rule):
    """通过XML直接设置段落间距，确保值被正确写入"""
    pPr = para._element.get_or_add_pPr()
    # 移除现有spacing
    existing = pPr.find(qn('w:spacing'))
    if existing is not None:
        pPr.remove(existing)
    spacing = OxmlElement('w:spacing')
    if before_twips is not None:
        spacing.set(qn('w:before'), str(before_twips))
    if after_twips is not None:
        spacing.set(qn('w:after'), str(after_twips))
    if line_twips is not None:
        spacing.set(qn('w:line'), str(line_twips))
    if line_rule is not None:
        spacing.set(qn('w:lineRule'), line_rule)
    pPr.append(spacing)


def set_para_indent_xml(para, first_line_twips=None, left_twips=None):
    """通过XML直接设置段落缩进"""
    pPr = para._element.get_or_add_pPr()
    existing = pPr.find(qn('w:ind'))
    if existing is not None:
        pPr.remove(existing)
    ind = OxmlElement('w:ind')
    if first_line_twips is not None:
        ind.set(qn('w:firstLine'), str(first_line_twips))
    if left_twips is not None:
        ind.set(qn('w:left'), str(left_twips))
    pPr.append(ind)


def set_page_break_before(para):
    """设置段前分页"""
    pPr = para._element.get_or_add_pPr()
    existing = pPr.find(qn('w:pageBreakBefore'))
    if existing is None:
        pbb = OxmlElement('w:pageBreakBefore')
        pPr.insert(0, pbb)


def get_cn_font(run):
    """获取run的中文字体名"""
    rPr = run._element.find(qn('w:rPr'))
    if rPr is not None:
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is not None:
            return rFonts.get(qn('w:eastAsia'), '')
    return ''


def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    """设置单元格边框
    参数: None=不设置, 'none'=无边框, (size_eighths, val)=有边框
    如 (12, 'single')=1.5磅, (8, 'single')=1磅
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    existing = tcPr.find(qn('w:tcBorders'))
    if existing is not None:
        tcPr.remove(existing)
    tcBorders = OxmlElement('w:tcBorders')
    for edge, spec in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if spec is None:
            continue
        border = OxmlElement(f'w:{edge}')
        if spec == 'none':
            border.set(qn('w:val'), 'nil')
        else:
            size_val, line_val = spec
            border.set(qn('w:val'), line_val)
            border.set(qn('w:sz'), str(size_val))
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
        tcBorders.append(border)
    tcPr.append(tcBorders)


# ============================================================
# 格式化函数
# ============================================================

def fix_page_setup(doc):
    """1. 页面设置：页边距3cm，页眉页脚2.2cm，A4"""
    for section in doc.sections:
        section.top_margin = Cm(3.0)
        section.bottom_margin = Cm(3.0)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(3.0)
        section.header_distance = Cm(2.2)
        section.footer_distance = Cm(2.2)
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
    print("[1] 页面设置: 3cm/2.2cm/A4")


def fix_style_colors(doc):
    """2. 将所有样式中的字体颜色改为黑色"""
    count = 0
    for style in doc.styles:
        if hasattr(style, 'element'):
            rPr = style.element.find(qn('w:rPr'))
            if rPr is not None:
                c = rPr.find(qn('w:color'))
                if c is not None:
                    val = c.get(qn('w:val'))
                    if val and val not in ('000000', 'auto', 'Auto'):
                        c.set(qn('w:val'), '000000')
                        count += 1
    # docDefaults
    styles_element = doc.styles.element
    docDefaults = styles_element.find(qn('w:docDefaults'))
    if docDefaults is not None:
        rPrDefault = docDefaults.find(qn('w:rPrDefault'))
        if rPrDefault is not None:
            rPr = rPrDefault.find(qn('w:rPr'))
            if rPr is not None:
                c = rPr.find(qn('w:color'))
                if c is None:
                    c = OxmlElement('w:color')
                    rPr.append(c)
                c.set(qn('w:val'), '000000')
    print(f"[2] 样式颜色: {count}个样式改为黑色")


def fix_all_run_colors(doc):
    """3. 显式设置所有run颜色为黑色"""
    count = 0
    for para in doc.paragraphs:
        for run in para.runs:
            if run.text.strip():
                run.font.color.rgb = RGBColor(0, 0, 0)
                count += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.text.strip():
                            run.font.color.rgb = RGBColor(0, 0, 0)
                            count += 1
    print(f"[3] Run颜色: {count}个run设为黑色")


def fix_headings(doc):
    """4. 修正章标题和节标题格式"""
    cn_num_map = {"一":1,"二":2,"三":3,"四":4,"五":5,"六":6,"七":7,"八":8,"九":9,"十":10}
    h1_count = 0
    h2_count = 0

    for i, para in enumerate(doc.paragraphs):
        style_name = para.style.name if para.style else ""
        text = para.text.strip()
        if not text:
            continue

        if style_name == "Heading 1":
            # 章标题命名格式转换："一、问题" → "第1章　问题"
            m = re.match(r'^([一二三四五六七八九十])、\s*(.+)$', text)
            if m:
                cn_num = m.group(1)
                title = m.group(2)
                if cn_num in cn_num_map:
                    new_text = f"第{cn_num_map[cn_num]}章\u3000{title}"
                    for run in para.runs:
                        run.text = ""
                    if para.runs:
                        para.runs[0].text = new_text

            # 格式：16pt黑体/Arial居中，单倍行距，段前24磅段后18磅
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_para_spacing_xml(para, 480, 360, 240, 'auto')
            set_para_indent_xml(para, first_line_twips=0)
            set_page_break_before(para)
            for run in para.runs:
                if run.text.strip():
                    set_run_font(run, "黑体", "Arial", 16, bold=False, italic=False)
            h1_count += 1

        elif style_name == "Heading 2":
            # 14pt黑体/Arial居左，固定值20磅，段前24磅段后6磅
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_para_spacing_xml(para, 480, 120, 400, 'exact')
            set_para_indent_xml(para, first_line_twips=0)
            for run in para.runs:
                if run.text.strip():
                    set_run_font(run, "黑体", "Arial", 14, bold=False, italic=False)
            h2_count += 1

        elif style_name == "Heading 3":
            # 13pt黑体/Arial居左，固定值20磅，段前12磅段后6磅
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_para_spacing_xml(para, 240, 120, 400, 'exact')
            set_para_indent_xml(para, first_line_twips=0)
            for run in para.runs:
                if run.text.strip():
                    set_run_font(run, "黑体", "Arial", 13, bold=False, italic=False)

    print(f"[4] 标题: H1={h1_count}, H2={h2_count}")


def fix_body_text(doc, skip_indices=None):
    """5. 修正正文段落格式"""
    skip_indices = skip_indices or set()
    count = 0
    for i, para in enumerate(doc.paragraphs):
        if i in skip_indices:
            continue
        style_name = para.style.name if para.style else ""
        text = para.text.strip()
        if style_name != "Normal" or not text or i <= 8:
            continue
        # 跳过图题、表题、参考文献
        if text.startswith("图") and re.match(r"^图\d", text):
            continue
        if re.match(r"^表\d", text):
            continue
        if re.match(r"^\[\d+\]", text):
            continue

        # 12pt宋体/TNR，两端对齐，首行缩进2字符，固定值20磅，段前0段后0
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_para_spacing_xml(para, 0, 0, 400, 'exact')
        set_para_indent_xml(para, first_line_twips=480)
        for run in para.runs:
            if run.text.strip():
                was_bold = run.font.bold
                set_run_font(run, "宋体", "Times New Roman", 12, bold=was_bold or False, italic=False)
        count += 1
    print(f"[5] 正文: {count}个段落")


def fix_figure_captions(doc, caption_map=None):
    """6. 修正图题格式
    caption_map: 可选的 {关键词: {"title": "图X.Y　标题", }} 映射
    如不提供，仅格式化已有的"图X.Y"开头的段落
    """
    count = 0
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not (text.startswith("图") and re.match(r"^图\d+\.\d", text)):
            continue

        # 如果有映射，替换文本
        if caption_map:
            for keywords, cap_info in caption_map.items():
                if keywords in text:
                    for run in para.runs:
                        run.text = ""
                    if para.runs:
                        para.runs[0].text = cap_info["title"]
                    break

        # 11pt宋体/TNR居中，单倍行距，段前6磅段后12磅
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_spacing_xml(para, 120, 240, 240, 'auto')
        set_para_indent_xml(para, first_line_twips=0)
        for run in para.runs:
            if run.text.strip():
                set_run_font(run, "宋体", "Times New Roman", 11, bold=False, italic=False)
        count += 1

        # 删除紧随其后的来源段落
        if i + 1 < len(doc.paragraphs):
            next_text = doc.paragraphs[i + 1].text.strip()
            if next_text.startswith("（来源") or next_text.startswith("(来源"):
                p = doc.paragraphs[i + 1]._element
                p.getparent().remove(p)

    print(f"[6] 图题: {count}个")


def fix_table_captions(doc):
    """7. 修正表题格式"""
    count = 0
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not re.match(r"^表\d+\.\d", text):
            continue
        # 11pt宋体/TNR居中，单倍行距，段前12磅段后6磅
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_spacing_xml(para, 240, 120, 240, 'auto')
        set_para_indent_xml(para, first_line_twips=0)
        for run in para.runs:
            if run.text.strip():
                set_run_font(run, "宋体", "Times New Roman", 11, bold=False, italic=False)
        count += 1
    print(f"[7] 表题: {count}个")


def fix_three_line_tables(doc):
    """8. 修正三线表格式"""
    for ti, table in enumerate(doc.tables):
        rows = table.rows
        num_rows = len(rows)
        for row_idx, row in enumerate(rows):
            for cell in row.cells:
                top_border = None
                bottom_border = None
                if row_idx == 0:
                    top_border = (12, 'single')   # 1.5磅
                    bottom_border = (8, 'single')  # 1磅
                elif row_idx == num_rows - 1:
                    bottom_border = (12, 'single') # 1.5磅
                set_cell_border(cell, top=top_border, bottom=bottom_border,
                               left='none', right='none')
                # 单元格文字：11pt宋体/TNR居中，单倍行距，段前3磅段后3磅
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    set_para_spacing_xml(para, 60, 60, 240, 'auto')
                    set_para_indent_xml(para, first_line_twips=0)
                    is_header = (row_idx == 0)
                    for run in para.runs:
                        if run.text.strip():
                            set_run_font(run, "宋体", "Times New Roman", 11,
                                       bold=is_header, italic=False)
    print(f"[8] 三线表: {len(doc.tables)}个")


def fix_references(doc):
    """9. 修正参考文献格式"""
    count = 0
    for para in doc.paragraphs:
        text = para.text.strip()
        if not re.match(r"^\[\d+\]", text):
            continue
        # 10.5pt宋体/TNR，固定值16磅，段前3磅段后0，悬挂缩进2字符
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_para_spacing_xml(para, 60, 0, 320, 'exact')
        # 悬挂缩进：left=420, firstLine=-420
        pPr = para._element.get_or_add_pPr()
        existing = pPr.find(qn('w:ind'))
        if existing is not None:
            pPr.remove(existing)
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), '420')
        ind.set(qn('w:firstLine'), '-420')
        pPr.append(ind)
        for run in para.runs:
            if run.text.strip():
                set_run_font(run, "宋体", "Times New Roman", 10.5, bold=False, italic=False)
        count += 1
    print(f"[9] 参考文献: {count}条")


def fix_cover(doc, cover_indices=None):
    """10. 封面字体设置
    cover_indices: {index: (cn_font, en_font, size_pt, bold)} 映射
    """
    if cover_indices is None:
        # 默认：前9段中找非空段落
        cover_indices = {}
        for i in range(min(9, len(doc.paragraphs))):
            text = doc.paragraphs[i].text.strip()
            if not text:
                continue
            if i == 4 or (not cover_indices and len(text) > 10):
                cover_indices[i] = ("黑体", "Arial", 26, True)
            elif i == 5 or "报告" in text:
                cover_indices[i] = ("宋体", "Times New Roman", 16, False)
            else:
                cover_indices[i] = ("宋体", "Times New Roman", 12, False)

    for idx, (cn, en, sz, bold) in cover_indices.items():
        if idx < len(doc.paragraphs):
            para = doc.paragraphs[idx]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                if run.text.strip():
                    set_run_font(run, cn, en, sz, bold=bold, italic=False)
    print(f"[10] 封面: {len(cover_indices)}个段落")


def fix_header_footer(doc, header_text=None):
    """11. 页眉页脚设置"""
    for section in doc.sections:
        # 页眉
        header = section.header
        header.is_linked_to_previous = False
        for para in header.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                if run.text.strip():
                    set_run_font(run, "宋体", "Times New Roman", 10.5, bold=False, italic=False)

        # 页脚：添加PAGE域
        footer = section.footer
        footer.is_linked_to_previous = False
        for para in footer.paragraphs:
            for run in para.runs:
                run.text = ""
        if footer.paragraphs:
            para = footer.paragraphs[0]
        else:
            para = footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = para.add_run()
        rPr = run._element.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), '宋体')
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        rPr.append(rFonts)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '21')
        rPr.append(sz)

        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' PAGE '
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._element.append(fldChar1)
        run._element.append(instrText)
        run._element.append(fldChar2)

    print("[11] 页眉页脚: 页眉五号居中，页脚PAGE域五号居中")


def add_toc(doc):
    """12. 在摘要前插入目录（TOC域）"""
    abstract_idx = None
    for i, para in enumerate(doc.paragraphs):
        if para.style and para.style.name == "Heading 1" and para.text.strip() == "摘要":
            abstract_idx = i
            break

    if abstract_idx is None:
        print("[12] 目录: 未找到'摘要'段落，跳过")
        return

    abstract_elem = doc.paragraphs[abstract_idx]._element

    # 目录标题段落
    toc_title = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    pPr.append(jc)
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '480')
    spacing.set(qn('w:after'), '360')
    spacing.set(qn('w:line'), '240')
    spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)
    pbb = OxmlElement('w:pageBreakBefore')
    pPr.append(pbb)
    toc_title.append(pPr)

    run_elem = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), '黑体')
    rFonts.set(qn('w:ascii'), 'Arial')
    rFonts.set(qn('w:hAnsi'), 'Arial')
    rPr.append(rFonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '32')
    rPr.append(sz)
    run_elem.append(rPr)
    text_elem = OxmlElement('w:t')
    text_elem.text = '目录'
    text_elem.set(qn('xml:space'), 'preserve')
    run_elem.append(text_elem)
    toc_title.append(run_elem)

    abstract_elem.addprevious(toc_title)

    # TOC域段落
    toc_para = OxmlElement('w:p')
    pPr2 = OxmlElement('w:pPr')
    spacing2 = OxmlElement('w:spacing')
    spacing2.set(qn('w:line'), '400')
    spacing2.set(qn('w:lineRule'), 'exact')
    pPr2.append(spacing2)
    toc_para.append(pPr2)

    run2 = OxmlElement('w:r')
    fldBegin = OxmlElement('w:fldChar')
    fldBegin.set(qn('w:fldCharType'), 'begin')
    run2.append(fldBegin)
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' TOC \\o "1-2" \\h \\z \\u '
    run2.append(instrText)
    fldSep = OxmlElement('w:fldChar')
    fldSep.set(qn('w:fldCharType'), 'separate')
    run2.append(fldSep)

    hint_run = OxmlElement('w:r')
    rPr3 = OxmlElement('w:rPr')
    rFonts3 = OxmlElement('w:rFonts')
    rFonts3.set(qn('w:eastAsia'), '宋体')
    rFonts3.set(qn('w:ascii'), 'Times New Roman')
    rFonts3.set(qn('w:hAnsi'), 'Times New Roman')
    rPr3.append(rFonts3)
    sz3 = OxmlElement('w:sz')
    sz3.set(qn('w:val'), '24')
    rPr3.append(sz3)
    hint_run.append(rPr3)
    hint_text = OxmlElement('w:t')
    hint_text.text = '（请在Word中右键此处选择"更新域"以生成目录）'
    hint_run.append(hint_text)

    fldEnd = OxmlElement('w:fldChar')
    fldEnd.set(qn('w:fldCharType'), 'end')
    run2.append(fldEnd)

    toc_para.append(run2)
    toc_para.append(hint_run)
    toc_title.addnext(toc_para)

    # 摘要前分页
    abs_pPr = abstract_elem.find(qn('w:pPr'))
    if abs_pPr is not None:
        abs_pbb = abs_pPr.find(qn('w:pageBreakBefore'))
        if abs_pbb is None:
            abs_pbb = OxmlElement('w:pageBreakBefore')
            abs_pPr.insert(0, abs_pbb)

    print("[12] 目录: 已插入TOC域（打开Word后右键更新）")


def add_table_titles(doc, table_descriptions):
    """13. 为表格添加表题
    table_descriptions: list of (table_index, "表X.Y　标题") 
    需在fix_table_captions之前调用
    """
    for ti in range(len(doc.tables) - 1, -1, -1):
        table = doc.tables[ti]
        # 查找对应的表题
        title = None
        for idx, t in table_descriptions:
            if idx == ti:
                title = t
                break
        if not title:
            continue

        new_para = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'center')
        pPr.append(jc)
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:before'), '240')
        spacing.set(qn('w:after'), '120')
        spacing.set(qn('w:line'), '240')
        spacing.set(qn('w:lineRule'), 'auto')
        pPr.append(spacing)
        new_para.append(pPr)

        run_elem = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), '宋体')
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        rPr.append(rFonts)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '22')
        rPr.append(sz)
        run_elem.append(rPr)
        text_elem = OxmlElement('w:t')
        text_elem.text = title
        text_elem.set(qn('xml:space'), 'preserve')
        run_elem.append(text_elem)
        new_para.append(run_elem)

        table._element.addprevious(new_para)

    print(f"[13] 表题: 已为{len(table_descriptions)}个表格添加表题")


# ============================================================
# 主函数
# ============================================================

def format_thu_thesis(input_path, output_path=None):
    """一键格式化"""
    if output_path is None:
        output_path = input_path

    print("=" * 60)
    print(f"清华论文格式化: {input_path}")
    print("=" * 60)

    doc = Document(input_path)

    fix_page_setup(doc)
    fix_style_colors(doc)
    fix_all_run_colors(doc)
    fix_headings(doc)
    fix_body_text(doc)
    fix_figure_captions(doc)
    fix_table_captions(doc)
    fix_three_line_tables(doc)
    fix_references(doc)
    fix_header_footer(doc)
    add_toc(doc)

    doc.save(output_path)
    print(f"\n{'=' * 60}")
    print(f"完成！已保存至: {output_path}")
    print(f"{'=' * 60}")
    print("\n注意：打开Word后请右键目录处选择'更新域'以生成目录内容。")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("用法: python format_thu_thesis.py <input.docx> [output.docx]")
        sys.exit(1)
    input_path = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 else None
    format_thu_thesis(input_path, output_path)
