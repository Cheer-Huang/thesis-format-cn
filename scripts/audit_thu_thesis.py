#!/usr/bin/env python3
"""
清华大学研究生学位论文 DOCX 格式审计脚本
检查 Word 文档是否符合清华规范，输出详细报告。

用法:
    python audit_thu_thesis.py <input.docx>
"""

import sys
import re
from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH


def get_cn_font(run):
    rPr = run._element.find(qn('w:rPr'))
    if rPr is not None:
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is not None:
            return rFonts.get(qn('w:eastAsia'), '')
    return ''

def emu_to_pt(emu):
    if emu is None: return None
    return round(emu / 12700, 1)

def emu_to_cm(emu):
    if emu is None: return None
    return round(emu / 360000, 2)

def get_xml_spacing(para):
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None: return {}
    sp = pPr.find(qn('w:spacing'))
    if sp is None: return {}
    return {
        'before': sp.get(qn('w:before')),
        'after': sp.get(qn('w:after')),
        'line': sp.get(qn('w:line')),
        'lineRule': sp.get(qn('w:lineRule')),
    }

def check_color(run):
    rPr = run._element.find(qn('w:rPr'))
    if rPr is not None:
        c = rPr.find(qn('w:color'))
        if c is not None:
            val = c.get(qn('w:val'))
            if val and val not in ('000000', 'auto'):
                return val
    rgb = run.font.color.rgb if run.font.color else None
    if rgb and str(rgb) != '000000':
        return str(rgb)
    return None


def audit(input_path):
    doc = Document(input_path)
    issues = []

    print("=" * 60)
    print("清华大学学位论文格式审计")
    print("=" * 60)

    # 1. 页面设置
    print("\n--- 1. 页面设置 ---")
    for s in doc.sections:
        tm, bm = emu_to_cm(s.top_margin), emu_to_cm(s.bottom_margin)
        lm, rm = emu_to_cm(s.left_margin), emu_to_cm(s.right_margin)
        hd, fd = emu_to_cm(s.header_distance), emu_to_cm(s.footer_distance)
        pw, ph = emu_to_cm(s.page_width), emu_to_cm(s.page_height)
        print(f"  页边距: 上{tm} 下{bm} 左{lm} 右{rm}cm")
        print(f"  页眉距{hd}cm 页脚距{fd}cm  纸张:{pw}x{ph}cm")
        if tm != 3.0 or bm != 3.0 or lm != 3.0 or rm != 3.0:
            issues.append("页边距不是3cm")
        if hd != 2.2 or fd != 2.2:
            issues.append("页眉/页脚距边界不是2.2cm")

    # 2. 字体颜色
    print("\n--- 2. 字体颜色 ---")
    color_issues = 0
    for para in doc.paragraphs:
        for run in para.runs:
            if run.text.strip():
                c = check_color(run)
                if c:
                    color_issues += 1
                    if color_issues <= 3:
                        print(f"  非黑色: {c} | {run.text[:30]}")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.text.strip():
                            c = check_color(run)
                            if c:
                                color_issues += 1
    if color_issues == 0:
        print("  全部黑色 ✓")
    else:
        issues.append(f"发现{color_issues}处非黑色字体")

    # 3. 章标题
    print("\n--- 3. 章标题 (Heading 1) ---")
    for i, para in enumerate(doc.paragraphs):
        if para.style and para.style.name == "Heading 1":
            text = para.text.strip()
            sp = get_xml_spacing(para)
            cn, en, sz = "", "", None
            for r in para.runs:
                if r.text.strip():
                    cn, en = get_cn_font(r), r.font.name
                    sz = emu_to_pt(r.font.size) if r.font.size else None
                    break
            pissues = []
            if para.alignment != WD_ALIGN_PARAGRAPH.CENTER: pissues.append(f"对齐={para.alignment}")
            if cn != "黑体": pissues.append(f"中文={cn}")
            if en != "Arial": pissues.append(f"英文={en}")
            if sz != 16.0: pissues.append(f"字号={sz}pt")
            if sp.get('lineRule') != 'auto': pissues.append(f"行距规则={sp.get('lineRule')}")
            if sp.get('before') != '480': pissues.append(f"段前={sp.get('before')}")
            if sp.get('after') != '360': pissues.append(f"段后={sp.get('after')}")
            if not re.match(r'^第\d+章', text) and text not in ("摘要", "参考文献", "Abstract"):
                pissues.append("命名格式非'第X章'")
            status = "✓" if not pissues else "✗"
            print(f"  [{i:3d}] {status} \"{text}\"")
            if pissues:
                issues.extend([f"章标题[{i}]: {p}" for p in pissues])
                print(f"         {pissues}")

    # 4. 节标题
    print("\n--- 4. 节标题 (Heading 2) ---")
    ok, bad = 0, 0
    for i, para in enumerate(doc.paragraphs):
        if para.style and para.style.name == "Heading 2":
            text = para.text.strip()
            sp = get_xml_spacing(para)
            cn, en, sz = "", "", None
            for r in para.runs:
                if r.text.strip():
                    cn, en = get_cn_font(r), r.font.name
                    sz = emu_to_pt(r.font.size) if r.font.size else None
                    break
            pissues = []
            if para.alignment != WD_ALIGN_PARAGRAPH.LEFT: pissues.append("非居左")
            if cn != "黑体": pissues.append(f"中文={cn}")
            if en != "Arial": pissues.append(f"英文={en}")
            if sz != 14.0: pissues.append(f"字号={sz}pt")
            if sp.get('lineRule') != 'exact': pissues.append(f"行距规则={sp.get('lineRule')}")
            if sp.get('line') != '400': pissues.append(f"行距={sp.get('line')}")
            if sp.get('before') != '480': pissues.append(f"段前={sp.get('before')}")
            if sp.get('after') != '120': pissues.append(f"段后={sp.get('after')}")
            if pissues:
                bad += 1
                print(f"  [{i:3d}] ✗ \"{text[:35]}\" -> {pissues}")
                issues.extend([f"节标题[{i}]: {p}" for p in pissues])
            else:
                ok += 1
    print(f"  通过: {ok}, 有问题: {bad}")

    # 5. 正文
    print("\n--- 5. 正文段落 ---")
    ok, bad = 0, 0
    for i, para in enumerate(doc.paragraphs):
        if para.style and para.style.name == "Normal":
            text = para.text.strip()
            if not text or i <= 8: continue
            if text.startswith("图") and re.match(r"^图\d", text): continue
            if re.match(r"^表\d", text): continue
            if re.match(r"^\[\d+\]", text): continue
            if text in ("目录",): continue
            if "更新域" in text: continue

            sp = get_xml_spacing(para)
            cn, en, sz = "", "", None
            for r in para.runs:
                if r.text.strip():
                    cn, en = get_cn_font(r), r.font.name
                    sz = emu_to_pt(r.font.size) if r.font.size else None
                    break
            pissues = []
            if para.alignment != WD_ALIGN_PARAGRAPH.JUSTIFY: pissues.append(f"对齐={para.alignment}")
            if cn != "宋体": pissues.append(f"中文={cn}")
            if en != "Times New Roman": pissues.append(f"英文={en}")
            if sz != 12.0: pissues.append(f"字号={sz}pt")
            if sp.get('lineRule') != 'exact': pissues.append(f"行距规则={sp.get('lineRule')}")
            if sp.get('line') != '400': pissues.append(f"行距={sp.get('line')}")
            if sp.get('before') != '0': pissues.append(f"段前={sp.get('before')}")
            if sp.get('after') != '0': pissues.append(f"段后={sp.get('after')}")
            pPr = para._element.find(qn('w:pPr'))
            if pPr is not None:
                ind = pPr.find(qn('w:ind'))
                if ind is not None:
                    fl = ind.get(qn('w:firstLine'))
                    if fl and abs(int(fl) - 480) > 20:
                        pissues.append(f"首行缩进={fl}")
            if pissues:
                bad += 1
                if bad <= 3:
                    print(f"  [{i:3d}] ✗ \"{text[:30]}\" -> {pissues}")
            else:
                ok += 1
    print(f"  通过: {ok}, 有问题: {bad}")
    if bad:
        issues.append(f"正文段落有{bad}个问题")

    # 6. 图题
    print("\n--- 6. 图题 ---")
    ok, bad = 0, 0
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text.startswith("图") and re.match(r"^图\d+\.\d", text):
            sp = get_xml_spacing(para)
            cn, en, sz, bold, italic = "", "", None, None, None
            for r in para.runs:
                if r.text.strip():
                    cn, en = get_cn_font(r), r.font.name
                    sz = emu_to_pt(r.font.size) if r.font.size else None
                    bold, italic = r.font.bold, r.font.italic
                    break
            pissues = []
            if para.alignment != WD_ALIGN_PARAGRAPH.CENTER: pissues.append("非居中")
            if cn != "宋体": pissues.append(f"中文={cn}")
            if en != "Times New Roman": pissues.append(f"英文={en}")
            if sz != 11.0: pissues.append(f"字号={sz}pt")
            if bold: pissues.append("不应加粗")
            if italic: pissues.append("不应斜体")
            if sp.get('lineRule') != 'auto': pissues.append(f"行距规则={sp.get('lineRule')}")
            if sp.get('before') != '120': pissues.append(f"段前={sp.get('before')}")
            if sp.get('after') != '240': pissues.append(f"段后={sp.get('after')}")
            if pissues:
                bad += 1
                print(f"  [{i:3d}] ✗ \"{text[:35]}\" -> {pissues}")
            else:
                ok += 1
    print(f"  通过: {ok}, 有问题: {bad}")

    # 7. 表题
    print("\n--- 7. 表题 ---")
    ok, bad = 0, 0
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if re.match(r"^表\d+\.\d", text):
            sp = get_xml_spacing(para)
            cn, en, sz = "", "", None
            for r in para.runs:
                if r.text.strip():
                    cn, en = get_cn_font(r), r.font.name
                    sz = emu_to_pt(r.font.size) if r.font.size else None
                    break
            pissues = []
            if para.alignment != WD_ALIGN_PARAGRAPH.CENTER: pissues.append("非居中")
            if cn != "宋体": pissues.append(f"中文={cn}")
            if en != "Times New Roman": pissues.append(f"英文={en}")
            if sz != 11.0: pissues.append(f"字号={sz}pt")
            if sp.get('before') != '240': pissues.append(f"段前={sp.get('before')}")
            if sp.get('after') != '120': pissues.append(f"段后={sp.get('after')}")
            if pissues:
                bad += 1
                print(f"  [{i:3d}] ✗ \"{text[:35]}\" -> {pissues}")
            else:
                ok += 1
    print(f"  通过: {ok}, 有问题: {bad}")

    # 8. 三线表
    print("\n--- 8. 三线表边框 ---")
    ok, bad = 0, 0
    for ti, table in enumerate(doc.tables):
        rows = table.rows
        nr = len(rows)
        tbad = False
        for ri in range(nr):
            for cell in rows[ri].cells:
                tcPr = cell._tc.find(qn('w:tcPr'))
                if tcPr is None: continue
                tcB = tcPr.find(qn('w:tcBorders'))
                if tcB is None: continue
                top = tcB.find(qn('w:top'))
                bot = tcB.find(qn('w:bottom'))
                left = tcB.find(qn('w:left'))
                right = tcB.find(qn('w:right'))
                if ri == 0:
                    if top is None or top.get(qn('w:sz')) != '12': tbad = True
                    if bot is None or bot.get(qn('w:sz')) != '8': tbad = True
                elif ri == nr - 1:
                    if bot is None or bot.get(qn('w:sz')) != '12': tbad = True
                if left is not None and left.get(qn('w:val')) != 'nil': tbad = True
                if right is not None and right.get(qn('w:val')) != 'nil': tbad = True
        if tbad:
            bad += 1
            print(f"  Table {ti}: ✗")
        else:
            ok += 1
    print(f"  通过: {ok}, 有问题: {bad}")

    # 9. 参考文献
    print("\n--- 9. 参考文献 ---")
    ok, bad = 0, 0
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if re.match(r"^\[\d+\]", text):
            sp = get_xml_spacing(para)
            cn, en, sz = "", "", None
            for r in para.runs:
                if r.text.strip():
                    cn, en = get_cn_font(r), r.font.name
                    sz = emu_to_pt(r.font.size) if r.font.size else None
                    break
            pPr = para._element.find(qn('w:pPr'))
            ind = pPr.find(qn('w:ind')) if pPr is not None else None
            pissues = []
            if cn != "宋体": pissues.append(f"中文={cn}")
            if en != "Times New Roman": pissues.append(f"英文={en}")
            if sz != 10.5: pissues.append(f"字号={sz}pt")
            if sp.get('lineRule') != 'exact': pissues.append(f"行距规则={sp.get('lineRule')}")
            if sp.get('line') != '320': pissues.append(f"行距={sp.get('line')}")
            if sp.get('before') != '60': pissues.append(f"段前={sp.get('before')}")
            if sp.get('after') != '0': pissues.append(f"段后={sp.get('after')}")
            if ind is not None:
                fl = ind.get(qn('w:firstLine'))
                if fl and int(fl) >= 0: pissues.append(f"首行缩进={fl}(应为负)")
            if pissues:
                bad += 1
                if bad <= 2:
                    print(f"  [{i:3d}] ✗ \"{text[:40]}\" -> {pissues}")
            else:
                ok += 1
    print(f"  通过: {ok}, 有问题: {bad}")

    # 10. 其他
    print("\n--- 10. 其他检查 ---")
    # 分页符
    pb_count = 0
    for para in doc.paragraphs:
        pPr = para._element.find(qn('w:pPr'))
        if pPr is not None and pPr.find(qn('w:pageBreakBefore')) is not None:
            pb_count += 1
    print(f"  分页符: {pb_count}个")
    if pb_count == 0:
        issues.append("无分页符（规范要求每章另起页）")

    # 页脚
    footer_xml = ""
    for para in doc.sections[0].footer.paragraphs:
        footer_xml += para._element.xml
    has_page_field = 'PAGE' in footer_xml
    print(f"  页脚PAGE域: {'有 ✓' if has_page_field else '无 ✗'}")
    if not has_page_field:
        issues.append("页脚无页码")

    # 目录
    has_toc = any("目录" in p.text or "TOC" in p._element.xml for p in doc.paragraphs)
    print(f"  目录: {'有 ✓' if has_toc else '无 ✗'}")
    if not has_toc:
        issues.append("无目录")

    # 总结
    print("\n" + "=" * 60)
    print("审计总结")
    print("=" * 60)
    if not issues:
        print("  ✅ 全部格式符合规范！")
    else:
        print(f"  共发现 {len(issues)} 个问题：")
        for idx, iss in enumerate(issues, 1):
            print(f"  {idx}. {iss}")

    return len(issues)


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("用法: python audit_thu_thesis.py <input.docx>")
        sys.exit(1)
    audit(sys.argv[1])
